#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set default margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Define colors
DARK_GREEN = RGBColor(26, 71, 42)  # #1a472a
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)

# Add footer
section = sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Futdle GDD v0.1 — Confidencial | Página "
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.runs[0]
footer_run.font.size = Pt(9)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("FUTDLE")
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = DARK_GREEN

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Game Design Document v0.1")
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph()  # Spacing

# Helper function to add styled heading
def add_heading(text, level=1):
    if level == 1:
        h = doc.add_heading(text, level=1)
        h.style.font.color.rgb = BLACK
    else:
        h = doc.add_heading(text, level=2)
    return h

# Helper function to add bold field
def add_field(label, value):
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(value)

# Section 1: Visão Geral
add_heading("1. Visão Geral")
add_field("Nome", "Futdle")
add_field("Plataforma", "Android e iOS (Flutter)")
add_field("Gênero", "Quiz / Puzzle diário")
add_field("Público-alvo", "Fãs de futebol, principalmente brasileiro, com expansão internacional")
add_field("Conceito", "Jogo diário de adivinhação de clubes de futebol. O jogador tenta descobrir o time do dia digitando nomes de clubes — a cada tentativa, o jogo revela quais características do time digitado correspondem ao time correto.")

# Section 2: Conceito e Inspiração
add_heading("2. Conceito e Inspiração")
doc.add_paragraph("Inspirado no Wordle e no Loldle (loldle.net), o Futdle traz a mecânica de \"adivinhe com pistas\" para o universo do futebol. A cada tentativa, o jogador recebe feedback visual sobre atributos do time que tentou — direcionando-o ao time correto.")

# Section 3: Modos de Jogo
add_heading("3. Modos de Jogo")

add_heading("3.1 Modo Clássico (MVP)", 2)
doc.add_paragraph("O jogador digita o nome de qualquer time do banco de dados", style='List Bullet')
doc.add_paragraph("O jogo exibe uma linha com os atributos do time digitado, comparados ao time correto", style='List Bullet')
doc.add_paragraph("Verde = correto | Amarelo = parcial | Vermelho = errado", style='List Bullet')
doc.add_paragraph("Setas ↑↓ para atributos numéricos (ex: ano de fundação mais antigo = ↑)", style='List Bullet')
doc.add_paragraph("Sem limite de tentativas no modo livre; desafio diário com registro de score", style='List Bullet')

add_heading("3.2 Modo Escudo (v2)", 2)
doc.add_paragraph("O escudo do time aparece coberto/pixelado", style='List Bullet')
doc.add_paragraph("A cada tentativa errada, uma parte é revelada", style='List Bullet')
doc.add_paragraph("O jogador tenta adivinhar o time pela imagem progressiva", style='List Bullet')

add_heading("3.3 Desafio Diário", 2)
doc.add_paragraph("Um time fixo por dia, igual para todos os jogadores", style='List Bullet')
doc.add_paragraph("Sem custo de energia", style='List Bullet')
doc.add_paragraph("Resultado compartilhável (estilo Wordle)", style='List Bullet')
doc.add_paragraph("Ranking diário", style='List Bullet')

add_heading("3.4 Modo Livre", 2)
doc.add_paragraph("Joga quantas rodadas quiser", style='List Bullet')
doc.add_paragraph("Consome 1 energia por partida", style='List Bullet')
doc.add_paragraph("Rodadas ilimitadas enquanto tiver energia", style='List Bullet')

# Section 4: Atributos
add_heading("4. Atributos do Modo Clássico")

# Create table
table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
headers = ["Atributo", "Tipo", "Feedback"]
for i, header_text in enumerate(headers):
    header_cells[i].text = header_text
    # Style header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1a472a')
    header_cells[i]._element.get_or_add_tcPr().append(shading_elm)

# Data rows
data = [
    ["País", "Texto", "Verde/Vermelho"],
    ["Continente", "Texto", "Verde/Vermelho"],
    ["Liga", "Texto", "Verde/Amarelo/Vermelho"],
    ["Ano de fundação", "Número", "Verde + seta ↑↓"],
    ["Cor principal", "Cor", "Verde/Amarelo/Vermelho"],
    ["Cor secundária", "Cor", "Verde/Amarelo/Vermelho"],
    ["Títulos nacionais", "Número", "Verde + seta ↑↓"],
    ["Títulos internacionais", "Número", "Verde + seta ↑↓"],
]

for i, row_data in enumerate(data):
    row = table.rows[i + 1]
    for j, cell_text in enumerate(row_data):
        row.cells[j].text = cell_text

# Legend
doc.add_paragraph()
legend = doc.add_paragraph()
legend_run = legend.add_run("Legenda:")
legend_run.bold = True

doc.add_paragraph("Verde: igual ao time correto", style='List Bullet')
doc.add_paragraph("Amarelo: parcialmente correto (ex: mesma confederação, cor similar)", style='List Bullet')
doc.add_paragraph("Vermelho: diferente", style='List Bullet')
doc.add_paragraph("↑: o valor correto é maior", style='List Bullet')
doc.add_paragraph("↓: o valor correto é menor", style='List Bullet')

# Section 5: Banco de Dados
add_heading("5. Banco de Dados de Clubes (MVP)")
add_heading("Ligas incluídas na v1:", 2)

doc.add_paragraph("Premier League (Inglaterra) — 20 times", style='List Bullet')
doc.add_paragraph("La Liga (Espanha) — 20 times", style='List Bullet')
doc.add_paragraph("Bundesliga (Alemanha) — 18 times", style='List Bullet')
doc.add_paragraph("Serie A (Itália) — 20 times", style='List Bullet')
doc.add_paragraph("Ligue 1 (França) — 18 times", style='List Bullet')
doc.add_paragraph("Brasileirão Série A — 20 times", style='List Bullet')
doc.add_paragraph("Liga Profesional Argentina — 28 times", style='List Bullet')
doc.add_paragraph("Liga MX (México) — 18 times", style='List Bullet')

total = doc.add_paragraph()
total_run = total.add_run("Total estimado: ~160 times")
total_run.bold = True

add_heading("Estrutura de dados de cada clube:", 2)
doc.add_paragraph("id, nome, nome_alternativo, país, continente, liga, ano_fundação, cor_principal (hex), cor_secundária (hex), títulos_nacionais, títulos_internacionais, url_escudo")

# Section 6: Sistema de Energia
add_heading("6. Sistema de Energia")
doc.add_paragraph("Energia máxima: 5 vidas", style='List Bullet')
doc.add_paragraph("Regeneração: 1 vida a cada 30 minutos", style='List Bullet')
doc.add_paragraph("Custo por partida no modo livre: 1 energia", style='List Bullet')
doc.add_paragraph("Desafio diário: gratuito, não consome energia", style='List Bullet')

add_heading("Formas de recuperar energia:", 2)
doc.add_paragraph("Assistir anúncio (+1 energia, limitado a 3x por dia)", style='List Bullet')
doc.add_paragraph("Comprar pacote de energia (IAP)", style='List Bullet')
doc.add_paragraph("Aguardar regeneração", style='List Bullet')

# Section 7: Monetização
add_heading("7. Monetização")
add_heading("Ads", 2)
doc.add_paragraph("Banner discreto na tela principal", style='List Bullet')
doc.add_paragraph("Rewarded ad para ganhar energia extra", style='List Bullet')
doc.add_paragraph("Interstitial ocasional entre partidas no modo livre", style='List Bullet')

add_heading("In-App Purchases", 2)
doc.add_paragraph("Pacote Pequeno: 5 energias", style='List Bullet')
doc.add_paragraph("Pacote Médio: 15 energias + remoção de banner por 7 dias", style='List Bullet')
doc.add_paragraph("Pacote Grande: energia ilimitada por 30 dias", style='List Bullet')
doc.add_paragraph("Remove Ads: compra única permanente", style='List Bullet')

# Section 8: UX
add_heading("8. UX / Fluxo de Telas")
doc.add_paragraph("Splash / Onboarding")
doc.add_paragraph("Tela principal — escolha do modo (Clássico / Escudo / Desafio Diário)")
doc.add_paragraph("Tela de jogo — campo de busca + grid de tentativas")
doc.add_paragraph("Resultado — compartilhar / jogar novamente / ranking")
doc.add_paragraph("Loja — energia e IAPs")
doc.add_paragraph("Perfil — histórico, streak, conquistas")

# Section 9: Arquitetura Técnica
add_heading("9. Arquitetura Técnica")

add_heading("Mobile", 2)
doc.add_paragraph("Framework: Flutter (Dart)", style='List Bullet')
doc.add_paragraph("Gerenciamento de estado: Riverpod ou BLoC", style='List Bullet')
doc.add_paragraph("Navegação: GoRouter", style='List Bullet')

add_heading("Backend", 2)
doc.add_paragraph("Plataforma: Supabase (PostgreSQL + Auth + Storage + Realtime)", style='List Bullet')
doc.add_paragraph("API: REST via Supabase SDK para Flutter", style='List Bullet')
doc.add_paragraph("Armazenamento de escudos: Supabase Storage + CDN", style='List Bullet')

add_heading("Banco de Dados (Supabase)", 2)
doc.add_paragraph("Tabelas principais:")
doc.add_paragraph("clubs (id, name, country, continent, league_id, founded_year, primary_color, secondary_color, national_titles, international_titles, shield_url)", style='List Bullet')
doc.add_paragraph("leagues (id, name, country, continent)", style='List Bullet')
doc.add_paragraph("daily_challenges (id, date, club_id, mode)", style='List Bullet')
doc.add_paragraph("user_progress (user_id, challenge_id, attempts, solved, time_taken)", style='List Bullet')
doc.add_paragraph("user_energy (user_id, current_energy, last_regen_at)", style='List Bullet')

add_heading("Seed de Dados", 2)
doc.add_paragraph("Usar TheSportsDB API (gratuita) para importar dados base dos clubes", style='List Bullet')
doc.add_paragraph("Curadoria manual para cores, títulos e inconsistências", style='List Bullet')
doc.add_paragraph("Script Python para importação e limpeza dos dados", style='List Bullet')

# Section 10: Roadmap
add_heading("10. Roadmap")

add_heading("v1 — MVP", 2)
doc.add_paragraph("Modo Clássico com desafio diário", style='List Bullet')
doc.add_paragraph("8 ligas / ~160 times", style='List Bullet')
doc.add_paragraph("Sistema de energia", style='List Bullet')
doc.add_paragraph("Ads + IAPs básicos", style='List Bullet')
doc.add_paragraph("Android + iOS", style='List Bullet')

add_heading("v2", 2)
doc.add_paragraph("Modo Escudo (revelação progressiva)", style='List Bullet')
doc.add_paragraph("Mais ligas (Champions League pool, Copa Libertadores pool)", style='List Bullet')
doc.add_paragraph("Conquistas e streaks", style='List Bullet')
doc.add_paragraph("Melhorias de UX com base em feedback", style='List Bullet')

add_heading("v3", 2)
doc.add_paragraph("Multiplayer assíncrono (dois jogadores, mesmo desafio, quem acerta mais rápido/com menos tentativas vence)", style='List Bullet')
doc.add_paragraph("Ranking global e por amigos", style='List Bullet')
doc.add_paragraph("Temporadas semanais", style='List Bullet')

# Section 11: Considerações Finais
add_heading("11. Considerações Finais")
doc.add_paragraph("O Futdle tem potencial de viralização orgânica pelo mecanismo de compartilhamento do resultado diário (igual ao Wordle). O foco inicial no público brasileiro, com expansão internacional natural pelo uso de ligas globais, posiciona bem o produto para crescimento orgânico. A ausência de multiplayer no MVP reduz complexidade técnica e acelera o lançamento.")

# Save document
output_path = r"C:\Users\mathe\AppData\Roaming\Claude\local-agent-mode-sessions\71d8c341-6e45-4f2a-8425-d7a6ea165739\faf2a8bb-0170-4bfd-a844-6c30b7ce3f9e\agent\local_ditto_faf2a8bb-0170-4bfd-a844-6c30b7ce3f9e\outputs\Futdle_GDD.docx"
doc.save(output_path)
print(f"✓ Documento criado com sucesso em: {output_path}")
